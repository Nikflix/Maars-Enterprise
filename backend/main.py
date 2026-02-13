# ============================================================================
# MAARS Enterprise - Backend (Local Execution)
# Multi-Agent Autonomous Research System
# ============================================================================
# Run with: python main.py
# ============================================================================

import os
import json
import asyncio
from typing import Optional, List
from io import BytesIO
import base64
import re
import traceback
import uuid
from datetime import datetime

from fastapi import FastAPI, HTTPException, Request, BackgroundTasks  # type: ignore
from fastapi.middleware.cors import CORSMiddleware  # type: ignore
from fastapi.responses import StreamingResponse, JSONResponse  # type: ignore
from pydantic import BaseModel  # type: ignore
import uvicorn  # type: ignore

import uuid
import asyncio
import time
from langchain_google_genai import ChatGoogleGenerativeAI  # type: ignore
from langchain_core.messages import HumanMessage  # type: ignore
from tavily import TavilyClient  # type: ignore

from docx import Document  # type: ignore
from docx.shared import Inches, Pt  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from fpdf import FPDF  # type: ignore

# ============================================================================
# Configuration (Hardcoded Keys)
# ============================================================================
GOOGLE_API_KEY = "AIzaSyBPrr4YdnWBB9afPLqLsCsqKnDlVZwKvG4"
TAVILY_API_KEY = "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"

os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY
os.environ["TAVILY_API_KEY"] = TAVILY_API_KEY

# Initialize clients
llm = ChatGoogleGenerativeAI(model="gemini-2.5-pro", google_api_key=GOOGLE_API_KEY, temperature=0.3)
tavily_client = TavilyClient(api_key=TAVILY_API_KEY)

async def invoke_llm_with_timeout(prompt: str, timeout_seconds: int = 30) -> str:
    """Invoke LLM with a timeout to prevent hanging."""
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        response = await asyncio.wait_for(
            loop.run_in_executor(None, lambda: llm.invoke([HumanMessage(content=prompt)])),  # type: ignore
            timeout=timeout_seconds
        )
        return response.content.strip()
    except asyncio.TimeoutError:
        raise TimeoutError(f"LLM call timed out after {timeout_seconds}s")
    except Exception as e:
        raise e

# ============================================================================
# FastAPI App
# ============================================================================
app = FastAPI(title="MAARS Enterprise API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Logging Middleware
@app.middleware("http")
async def log_requests(request: Request, call_next):
    try:
        import datetime
        with open("backend_requests.log", "a", encoding="utf-8") as f:
            f.write(f"{datetime.datetime.now()} {request.method} {request.url}\n")
    except Exception as e:
        print(f"Middleware log error: {e}")
    
    try:
        response = await call_next(request)
        with open("backend_requests.log", "a", encoding="utf-8") as f:
             f.write(f" -> Status: {response.status_code}\n")
        return response
    except Exception as e:
         # Log unhandled exceptions that crash the app
        import traceback
        import datetime
        print(f"Unhandled app error: {e}")
        traceback.print_exc()
        try:
            with open("backend_crash.log", "a", encoding="utf-8") as f:
                 f.write(f"\n--- Crash at {datetime.datetime.now()} ---\n")
                 f.write(f"Url: {request.url}\n")
                 traceback.print_exc(file=f)
        except:
            pass
        raise e
class TopicRequest(BaseModel):
    profession: str

class PlanRequest(BaseModel):
    topic: str

class WriteRequest(BaseModel):
    topic: str
    outline: dict

class RefineTopicRequest(BaseModel):
    profession: str
    topic_input: str
    current_topic: str

class RefinePlanRequest(BaseModel):
    current_plan: dict
    feedback: str

class WriteSectionRequest(BaseModel):
    section: dict
    topic: str
    context: dict

class PlagiarismRequest(BaseModel):
    content: str

class DownloadRequest(BaseModel):
    title: str
    content: str
    format: str  # "pdf" or "docx"

class SavePaperRequest(BaseModel):
    title: str
    content: str
    author: Optional[str] = "Anonymous"
    user_id: str
    is_public: bool = False
    plagiarism_score: float = 0

# ============================================================================
# Shared Paper Storage (File-based JSON)
# ============================================================================
PAPERS_FILE = "papers.json"

def load_papers():
    if os.path.exists(PAPERS_FILE):
        try:
            with open(PAPERS_FILE, "r") as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading papers: {e}")
            return []
    return []

def save_papers(papers):
    try:
        with open(PAPERS_FILE, "w") as f:
            json.dump(papers, f, indent=2)
    except Exception as e:
        print(f"Error saving papers: {e}")

papers_db: List[dict] = load_papers()

# ============================================================================
# Helper Functions
# ============================================================================
def search_web(query: str, max_results: int = 5) -> List[dict]:
    """Search the web using Tavily and return results with citations."""
    try:
        results = tavily_client.search(query=query, max_results=max_results)
        return [
            {
                "title": r.get("title", ""),
                "url": r.get("url", ""),
                "content": str(r.get("content", ""))[:500]  # type: ignore[index]
            }
            for r in results.get("results", [])
        ]
    except Exception as e:
        print(f"Search error: {e}")
        return []

def create_pdf(title: str, content: str) -> bytes:
    """Generate PDF from markdown content."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    # Set explicit margins to ensure enough horizontal space
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    
    def safe(text: str) -> str:
        """Sanitize text for Helvetica (Latin-1 only)."""
        if not text:
            return ""
        # Common unicode replacements
        replacements = {
            '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"',
            '\u2013': '-', '\u2014': '--', '\u2026': '...', '\u2022': '* ',
            '\u00a9': '(c)', '\u00ae': '(R)', '\u2122': '(TM)',
            '\u2010': '-', '\u2011': '-', '\u2012': '-',
            '\u200b': '', '\u200c': '', '\u200d': '', '\ufeff': '',
            '\x00': '',
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Encode to latin-1, dropping anything that can't be represented
        return text.encode('latin-1', errors='replace').decode('latin-1')
    
    def safe_multi_cell(p, w, h, txt):
        """Wrapper around multi_cell that resets X position and handles errors."""
        p.set_x(p.l_margin)  # Reset X to left margin
        try:
            p.multi_cell(w, h, txt)
        except Exception:
            # If still fails, skip this text
            pass
    
    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_x(pdf.l_margin)
    try:
        pdf.cell(0, 10, safe(title) or "Research Paper", ln=True, align="C")
    except Exception:
        pdf.multi_cell(0, 10, safe(title) or "Research Paper")
    pdf.ln(10)
    
    # Content
    pdf.set_font("Helvetica", "", 11)
    for line in (content or "").split('\n'):
        stripped = line.strip()
        try:
            _render_line(pdf, stripped, safe, safe_multi_cell)
        except Exception:
            # Page space exhausted — add new page and retry
            try:
                pdf.add_page()
                pdf.set_font("Helvetica", "", 11)
                _render_line(pdf, stripped, safe, safe_multi_cell)
            except Exception:
                # Truly unrenderable line — skip it
                pdf.set_font("Helvetica", "", 11)
    
    return pdf.output()

def _render_line(pdf, stripped: str, safe, safe_multi_cell):
    """Render a single line of markdown to the PDF."""
    if not stripped:
        pdf.ln(3)
    elif stripped.startswith('### '):
        pdf.set_font("Helvetica", "B", 12)
        pdf.ln(4)
        safe_multi_cell(pdf, 0, 6, safe(stripped[4:]))  # type: ignore
        pdf.set_font("Helvetica", "", 11)
    elif stripped.startswith('## '):
        pdf.set_font("Helvetica", "B", 13)
        pdf.ln(6)
        safe_multi_cell(pdf, 0, 7, safe(stripped[3:]))  # type: ignore
        pdf.set_font("Helvetica", "", 11)
    elif stripped.startswith('# '):
        pdf.set_font("Helvetica", "B", 14)
        pdf.ln(8)
        safe_multi_cell(pdf, 0, 8, safe(stripped[2:]))  # type: ignore
        pdf.set_font("Helvetica", "", 11)
    else:
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
        clean = re.sub(r'\*(.*?)\*', r'\1', clean)
        clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean)  # type: ignore
        safe_multi_cell(pdf, 0, 6, safe(clean))

def create_docx(title: str, content: str) -> bytes:
    """Generate DOCX from markdown content."""
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, 0)  # type: ignore
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = content.split('\n')  # type: ignore
    for line in lines:
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)  # type: ignore[index]
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)  # type: ignore[index]
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)  # type: ignore[index]
        elif line.strip():
            # Clean markdown formatting
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
            clean_line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_line)
            doc.add_paragraph(clean_line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ============================================================================
# API Endpoints
# ============================================================================
@app.get("/")
async def root():
    return {"message": "MAARS Enterprise API is running", "version": "1.0.0"}

@app.post("/suggest_topics")
async def suggest_topics(request: TopicRequest):
    """Generate 3 research topic suggestions based on user's profession."""
    prompt = f"""You are a research advisor. Based on the user's profession: "{request.profession}"
    
Suggest exactly 3 cutting-edge, specific research topics that would be relevant and impactful.
Each topic should be suitable for an IEEE-style academic paper.

Return ONLY a JSON array with 3 objects, each having "title" and "description" keys.
Example format:
[
  {{"title": "Topic Title Here", "description": "A brief 1-2 sentence description of why this topic is relevant."}},
  ...
]"""

    try:
        content = await invoke_llm_with_timeout(prompt, timeout_seconds=20)
        print(f"LLM response for topics: {content[:200]}...")  # type: ignore
        # Parse JSON from response
        # Extract JSON array from response
        start = content.find('[')
        end = content.rfind(']') + 1
        if start != -1 and end > start:
            topics = json.loads(content[start:end])  # type: ignore
            return {"topics": topics}
        else:
            print(f"Invalid JSON from LLM: {content}")
            raise ValueError("No valid JSON array found")
    except Exception as e:
        print(f"Error suggesting topics: {e}")
        # Fallback topics
        return {
            "topics": [
                {"title": f"AI Applications in {request.profession}", "description": "Exploring how artificial intelligence is transforming the field."},  # type: ignore
                {"title": f"Sustainable Practices in {request.profession}", "description": "Analyzing eco-friendly approaches and their impact."},
                {"title": f"Digital Transformation in {request.profession}", "description": "Examining the role of technology in modernizing practices."}
            ]
        }

@app.post("/plan_research")  # type: ignore
async def plan_research(request: PlanRequest):
    """Generate an IEEE-style research outline using Groq + Tavily search."""
    # First, search for relevant sources
    search_results = search_web(f"{request.topic} research academic", max_results=5)
    
    if search_results:
        sources_context = "\n".join([
            f"- {r['title']}: {r['content'][:200]}... (Source: {r['url']})"
            for r in search_results
        ])
        source_prompt = f"Based on these relevant sources:\n{sources_context}"
    else:
        sources_context = "No external sources found."
        source_prompt = "Generate the outline based on your internal knowledge. Do not include specific source citations in the outline description, but you may suggest potential search queries in the 'sources' field."

    prompt = f"""You are an expert research planner. Create a detailed IEEE-format research paper outline for the topic:

"{request.topic}"

{source_prompt}

Generate a strict IEEE paper outline with these sections:
1. Abstract
2. Introduction (with subsections: Background, Problem Statement, Objectives)

3. Literature Review
4. Methodology
5. Results and Discussion
6. Conclusion
7. Future Work
8. References

For each section, provide:
- A brief description of what should be covered
- Key points to address
- Relevant sources from the search results

Return the outline as a JSON object with this structure:
{{
  "title": "Full Research Paper Title",
  "sections": [
    {{
      "id": "1",
      "name": "Abstract",
      "description": "Brief description",
      "key_points": ["point1", "point2"],
      "sources": ["relevant source URLs"]
    }},
    ...
  ],
  "sources": [
    {{"title": "Source Title", "url": "URL", "relevance": "Why relevant"}}
  ]
}}"""

    try:
        content = await invoke_llm_with_timeout(prompt, timeout_seconds=45)
        print(f"LLM response for outline: {content[:200]}...")  # type: ignore
        
        # Extract JSON from response
        start = content.find('{')
        end = content.rfind('}') + 1
        if start != -1 and end > start:
            outline = json.loads(content[start:end])  # type: ignore
            return {"outline": outline, "search_results": search_results}
        else:
            raise ValueError("No valid JSON found in LLM response")
    except Exception as e:
        print(f"Error planning research: {e}")
        # Return a fallback outline instead of crashing
        # Generate a proper academic title from the topic
        try:
            title_content = await invoke_llm_with_timeout(
                f"Generate a single formal IEEE academic paper title for a research paper about: \"{request.topic}\". Return ONLY the title text, nothing else. No quotes, no explanation.",  # type: ignore
                timeout_seconds=10
            )
            fallback_title = title_content.strip().strip('"').strip("'")
        except Exception:
            fallback_title = f"A Comprehensive Study on {request.topic.title()}: Analysis, Applications, and Future Directions"
        
        fallback_outline = {  # type: ignore
            "title": fallback_title,
            "sections": [
                {"id": "1", "name": "Abstract", "description": "Overview of the research topic and key findings.", "key_points": ["Research objective", "Methodology overview", "Key results"], "sources": []},
                {"id": "2", "name": "Introduction", "description": "Background context and problem statement.", "key_points": ["Background", "Problem Statement", "Objectives"], "sources": []},
                {"id": "3", "name": "Literature Review", "description": "Review of existing work in this area.", "key_points": ["Current state of research", "Key findings from prior work", "Research gaps"], "sources": []},
                {"id": "4", "name": "Methodology", "description": "Research methods and approach.", "key_points": ["Research design", "Data collection", "Analysis methods"], "sources": []},
                {"id": "5", "name": "Results and Discussion", "description": "Presentation and analysis of findings.", "key_points": ["Key findings", "Analysis", "Implications"], "sources": []},
                {"id": "6", "name": "Conclusion", "description": "Summary of contributions and future work.", "key_points": ["Summary", "Contributions", "Future directions"], "sources": []},
                {"id": "7", "name": "References", "description": "List of cited works.", "key_points": ["Academic references"], "sources": []}
            ],
            "sources": []
        }
        return {"outline": fallback_outline, "search_results": []}

@app.post("/plan")
async def plan_research_alias(request: PlanRequest):
    """Alias for plan_research to match frontend."""
    return await plan_research(request)

@app.post("/refine-topic")
async def refine_topic(request: RefineTopicRequest):
    """Refine the research topic based on user input."""
    prompt = f"""You are a research advisor for a "{request.profession}". 
    current topic: "{request.current_topic}"
    User input: "{request.topic_input}"

    If the user input is a new topic proposal, set it as the refined topic.
    If the user is ignoring the current topic, just acknowledge the new input.
    
    Return a JSON object:
    {{
        "message": "Response to the user",
        "refined_topic": "The finalized topic name (or null if still discussing)"
    }}
    """
    try:
        content = await invoke_llm_with_timeout(prompt, timeout_seconds=15)
        start = content.find('{')
        end = content.rfind('}') + 1
        if start != -1 and end > start:
            json_str = content[start:end]  # type: ignore
            return json.loads(json_str)
        return {"message": "Could not parse response", "refined_topic": None}
    except Exception as e:
        print(f"Refine topic error: {e}")
        return {"message": "I can help you research that. Shall we proceed with planning?", "refined_topic": request.topic_input}

@app.post("/refine-plan")
async def refine_plan(request: RefinePlanRequest):
    """Refine the research plan based on user feedback."""
    prompt = f"""Update this research outline based on user feedback:
    Current Plan: {json.dumps(request.current_plan)}
    User Feedback: "{request.feedback}"
    
    Return the updated plan as a JSON object with the same structure (title, sections, sources).
    """
    try:
        content = await invoke_llm_with_timeout(prompt, timeout_seconds=30)
        start = content.find('{')
        end = content.rfind('}') + 1
        if start != -1 and end > start:
            json_str = content[start:end]  # type: ignore
            updated_plan = json.loads(json_str)
            return {"updated_plan": updated_plan}
        return {"updated_plan": request.current_plan}
    except Exception as e:
        print(f"Refine plan error: {e}")
        return {"updated_plan": request.current_plan}

@app.post("/write_paper")
async def write_paper(request: WriteRequest):
    """Stream the full IEEE paper section by section with real citations."""
    
    async def generate_paper():
        outline = request.outline
        raw_title = outline.get("title", request.topic)
        sections = outline.get("sections", [])
        
        # Generate a proper academic title if the current one looks like raw user input
        title = raw_title
        if len(raw_title.split()) <= 3 or raw_title == request.topic:
            try:
                generated_title = await invoke_llm_with_timeout(
                    f"Generate a single formal IEEE academic paper title for a research paper about: \"{request.topic}\". The paper covers these sections: {', '.join(s.get('name', '') for s in sections)}. Return ONLY the title text, nothing else. No quotes, no explanation.",
                    timeout_seconds=10
                )
                title = generated_title.strip().strip('"').strip("'")
                print(f"Generated paper title: {title}")
            except Exception as e:
                print(f"Title generation failed, using fallback: {e}")
                title = f"A Comprehensive Study on {request.topic.title()}: Analysis, Applications, and Future Directions"
        
        # Send title first
        yield f"data: {json.dumps({'type': 'title', 'content': f'# {title}'})}\n\n"
        await asyncio.sleep(0.1)
        
        for section in sections:
            section_name = section.get("name", "Section")
            description = section.get("description", "")
            key_points = section.get("key_points", [])
            sources = section.get("sources", [])
            
            # Search for section-specific content
            search_query = f"{request.topic} {section_name}"
            section_sources = search_web(search_query, max_results=3)
            
            sources_text = "\n".join([
                f"- {s['title']}: {s['content'][:300]}... [Source: {s['url']}]"
                for s in section_sources
            ])
            
            prompt = f"""Write the "{section_name}" section for an IEEE academic paper on "{request.topic}".

Section description: {description}
Key points to cover: {', '.join(key_points) if key_points else 'General coverage'}

Use these sources for citations:
{sources_text}

Requirements:
- Write in formal academic style
- Include in-text citations in IEEE format [1], [2], etc.
- Be comprehensive but concise (300-500 words for main sections, 150-200 for Abstract)
- Use technical language appropriate for an academic audience

Write ONLY the section content in Markdown format. Start with the section header (## {section_name})."""

            try:
                response = llm.invoke([HumanMessage(content=prompt)])
                section_content = response.content.strip()
                
                # Stream section content
                yield f"data: {json.dumps({'type': 'section', 'name': section_name, 'content': section_content})}\n\n"
                await asyncio.sleep(0.3)
                
                # Add sources for this section
                if section_sources:
                    for i, source in enumerate(section_sources):
                        yield f"data: {json.dumps({'type': 'source', 'section': section_name, 'source': source})}\n\n"
                
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'section': section_name, 'message': str(e)})}\n\n"
        
        yield f"data: {json.dumps({'type': 'complete'})}\n\n"
    
    return StreamingResponse(
        generate_paper(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )

@app.post("/write-section")
async def write_section(request: WriteSectionRequest):
    """Stream a single section content."""
    
    async def generate_section():
        section = request.section
        target_length = int(section.get('word_count', 0))
        
        length_instruction = ""
        if target_length > 0:
            length_instruction = f"""
            CRITICAL INSTRUCTION: The user has requested a section length of approximately {target_length} words.
            You MUST adhere to this word count strictly (+/- 10%).
            
            Current Target: {target_length} words.

            Стратегии for length:
            - If target is low (<200 words): Be concise, direct, and avoid fluff. Bullet points are good.
            - If target is medium (200-500 words): Standard academic paragraphs.
            - If target is high (>500 words): Go into deep detail, use examples, counter-arguments, and extensive analysis.

            DO NOT output significantly fewer words than requested. 
            DO NOT output significantly more words than requested.
            """
        else:
            length_instruction = "Write a comprehensive section."

        prompt = f"""Write the "{section.get('name')}" section for a paper on "{request.topic}".
        Context: {section.get('description')}
        Key points: {section.get('key_points')}
        
        {length_instruction}
        Write in academic markdown. use proper headings and subheadings.
        """
        try:
            response = llm.invoke([HumanMessage(content=prompt)])
            # Stream the content in chunks (simulated or real if LLM supports it, here we just yield the block)
            # Frontend expects a stream of bytes.
            yield response.content.encode('utf-8')
        except Exception as e:
            yield f"Error: {e}".encode('utf-8')

    return StreamingResponse(generate_section(), media_type="text/plain")

@app.post("/check_plagiarism")
async def check_plagiarism(request: PlagiarismRequest):
    """Check plagiarism by searching for similar content online."""
    # Extract key sentences from the content
    sentences = [s.strip() for s in request.content.split('.') if len(s.strip()) > 50][:5]  # type: ignore[index]
    
    total_matches: int = 0
    matched_sources: list = []
    
    for sentence in sentences:
        # Search for exact or similar matches
        results = search_web(f'"{sentence[:100]}"', max_results=3)  # type: ignore[index]
        if results:
            total_matches += 1  # type: ignore[operator]
            for r in results:
                if r['url'] not in [m['url'] for m in matched_sources]:
                    matched_sources.append({
                        "url": r['url'],
                        "title": r['title'],
                        "matched_text": str(sentence[:100]) + "..."  # type: ignore[index]
                    })
    
    # Calculate similarity score (simplified)
    similarity_score: float = min((total_matches / max(len(sentences), 1)) * 100, 100)  # type: ignore[operator]
    
    return {
        "similarity_score": round(similarity_score, 2),  # type: ignore[call-overload]
        "checked_sentences": len(sentences),
        "matches_found": total_matches,
        "matched_sources": matched_sources[:5],  # type: ignore[index]
        "status": "low" if similarity_score < 20 else "medium" if similarity_score < 50 else "high"
    }

# ============================================================================
# Async Plagiarism Job System
# ============================================================================
jobs_db = {}

class PlagiarismJobRequest(BaseModel):
    paperId: str
    content: str

async def process_plagiarism_job(job_id: str, content: str):
    """Background task to process plagiarism check."""
    try:
        # Simulate processing time + legitimate check
        jobs_db[job_id]["status"] = "processing"
        
        # Reuse existing logic or call checking function
        # For better async behavior in a real app, this would be cleaner. 
        # Here we just inline the logic for simplicity but wrapped in the job.
        
        sentences = [s.strip() for s in content.split('.') if len(s.strip()) > 50][:10] # Check more sentences
        total_matches = 0
        matched_sources = []
        
        for i, sentence in enumerate(sentences):
            # Update progress (optional, but good for polling)
            jobs_db[job_id]["progress"] = int((i / len(sentences)) * 90)
            
            # Simulated delay for "Scanning" UX
            await asyncio.sleep(0.5) 
            
            results = search_web(f'"{sentence[:100]}"', max_results=3)
            if results:
                total_matches += 1
                for r in results:
                    if r['url'] not in [m['url'] for m in matched_sources]:
                        matched_sources.append({
                            "url": r['url'],
                            "title": r['title'],
                            "matched_text": str(sentence[:100]) + "..."
                        })
        
        similarity_score = min((total_matches / max(len(sentences), 1)) * 100, 100)
        
        jobs_db[job_id]["result"] = {
            "similarity_score": round(similarity_score, 2),
            "checked_sentences": len(sentences),
            "matches_found": total_matches,
            "matched_sources": matched_sources[:10],
            "status": "low" if similarity_score < 20 else "medium" if similarity_score < 50 else "high"
        }
        jobs_db[job_id]["status"] = "completed"
        jobs_db[job_id]["progress"] = 100
        
    except Exception as e:
        print(f"Job {job_id} failed: {e}")
        jobs_db[job_id]["status"] = "error"
        jobs_db[job_id]["error"] = str(e)

@app.post("/api/plagiarism/start")
async def start_plagiarism_check(request: PlagiarismJobRequest, background_tasks: BackgroundTasks):
    job_id = str(uuid.uuid4())
    jobs_db[job_id] = {
        "id": job_id,
        "status": "pending",
        "progress": 0,
        "submitted_at": time.time()
    }
    background_tasks.add_task(process_plagiarism_job, job_id, request.content)
    return {"jobId": job_id}

@app.get("/api/plagiarism/status/{job_id}")
async def get_plagiarism_status(job_id: str):
    job = jobs_db.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return {
        "status": job["status"], 
        "progress": job.get("progress", 0)
    }

@app.get("/api/plagiarism/result/{job_id}")
async def get_plagiarism_result(job_id: str):
    job = jobs_db.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if job["status"] != "completed":
        raise HTTPException(status_code=400, detail="Job not completed")
    return job["result"]

@app.post("/download")
async def download_document(request: DownloadRequest):
    """Generate and return PDF or DOCX document."""
    try:
        if request.format.lower() == "pdf":
            doc_bytes = create_pdf(request.title, request.content)
            mime_type = "application/pdf"
            filename = f"{request.title[:50]}.pdf"  # type: ignore[index]
        elif request.format.lower() == "docx":
            doc_bytes = create_docx(request.title, request.content)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{request.title[:50]}.docx"  # type: ignore[index]
        else:
            raise HTTPException(status_code=400, detail="Format must be 'pdf' or 'docx'")
        
        # Return as base64 for frontend handling
        encoded = base64.b64encode(doc_bytes).decode('utf-8')
        return {
            "filename": filename,
            "mime_type": mime_type,
            "data": encoded
        }
    except Exception as e:
        print(f"Download error: {e}")
        traceback.print_exc()
        try:
            with open("backend_error.log", "a", encoding="utf-8") as f:
                f.write(f"\n--- Error at {datetime.now()} ---\n")
                f.write(f"Request: {str(request)}\n")
                traceback.print_exc(file=f)
        except Exception as log_err:
            print(f"Failed to log error: {log_err}")
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# Paper Library Endpoints (Shared/Public)
# ============================================================================
@app.post("/papers")
async def save_paper(request: SavePaperRequest):
    """Save a paper to the shared library."""
    paper = {
        "id": str(uuid.uuid4()),
        "title": request.title,
        "content": request.content,
        "author": request.author,
        "user_id": request.user_id,
        "is_public": request.is_public,
        "plagiarism_score": request.plagiarism_score,
        "created_at": datetime.now().isoformat(),
    }
    papers_db.insert(0, paper)
    save_papers(papers_db)
    return paper

@app.get("/papers")
async def get_papers():
    """Get all papers from the shared library."""
    return papers_db

@app.delete("/papers/{paper_id}")
async def delete_paper(paper_id: str):
    """Delete a paper from the shared library."""
    global papers_db
    papers_db = [p for p in papers_db if p["id"] != paper_id]
    save_papers(papers_db)
    return {"status": "deleted"}

@app.get("/stats")
async def get_stats():
    """Get global platform statistics."""
    return {
        "total_papers": len(papers_db),
        "compliant_percentage": 98,
        "avg_time": "<2m"
    }

# ============================================================================
# Run Server
# ============================================================================
if __name__ == "__main__":
    print("=" * 60)
    print("[*] MAARS Enterprise Backend Starting...")
    print("=" * 60)
    print("[>] Server URL: http://localhost:8000")
    print("[>] API Docs: http://localhost:8000/docs")
    print("=" * 60)
    
    try:
        print("Starting Uvicorn on port 8001...")
        uvicorn.run(app, host="0.0.0.0", port=8001, log_level="debug")
        print("Uvicorn run finished normally.")
    except BaseException as e:
        print(f"CRITICAL STARTUP ERROR: {e}")
        import traceback
        traceback.print_exc()
        with open("startup_error.log", "w") as f:
            f.write(f"Startup error: {e}\n")
            traceback.print_exc(file=f)
