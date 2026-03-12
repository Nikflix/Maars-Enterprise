import re
import base64
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

class DocumentService:
    def _render_line(self, pdf, stripped: str, safe, safe_multi_cell):
        """Render a single line of markdown to the PDF."""
        if not stripped:
            pdf.ln(3)
        # Skip code fence lines
        elif stripped.startswith('```'):
            return
        elif stripped.startswith('### '):
            pdf.set_font("Helvetica", "B", 12)
            pdf.ln(4)
            safe_multi_cell(pdf, 0, 6, safe(stripped[4:]))
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith('## '):
            pdf.set_font("Helvetica", "B", 13)
            pdf.ln(6)
            safe_multi_cell(pdf, 0, 7, safe(stripped[3:]))
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith('# '):
            pdf.set_font("Helvetica", "B", 14)
            pdf.ln(8)
            safe_multi_cell(pdf, 0, 8, safe(stripped[2:]))
            pdf.set_font("Helvetica", "", 11)
        else:
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            clean = re.sub(r'\*(.*?)\*', r'\1', clean)
            clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean)
            safe_multi_cell(pdf, 0, 6, safe(clean))

    def create_pdf(self, title: str, content: str) -> bytes:
        """Generate PDF from markdown content."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        # Set explicit margins to ensure enough horizontal space
        pdf.set_left_margin(15)
        pdf.set_right_margin(15)
        
        def safe(text: str) -> str:
            """Sanitize text for Helvetica (Latin-1 only)."""
            if not text:
                return ""
            # Common unicode replacements
            replacements = {
                '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"',
                '\u2013': '-', '\u2014': '--', '\u2026': '...', '\u2022': '* ',
                '\u00a9': '(c)', '\u00ae': '(R)', '\u2122': '(TM)',
                '\u2010': '-', '\u2011': '-', '\u2012': '-',
                '\u200b': '', '\u200c': '', '\u200d': '', '\ufeff': '',
                '\x00': '',
            }
            for old, new in replacements.items():
                text = text.replace(old, new)
            # Encode to latin-1, dropping anything that can't be represented
            return text.encode('latin-1', errors='replace').decode('latin-1')
        
        def safe_multi_cell(p, w, h, txt):
            """Wrapper around multi_cell that resets X position and handles errors."""
            p.set_x(p.l_margin)  # Reset X to left margin
            try:
                p.multi_cell(w, h, txt)
            except Exception:
                # If still fails, skip this text
                pass
        
        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_x(pdf.l_margin)
        try:
            pdf.cell(0, 10, safe(title) or "Research Paper", ln=True, align="C")
        except Exception:
            pdf.multi_cell(0, 10, safe(title) or "Research Paper")
        pdf.ln(10)
        
        # Content - strip any remaining code fences
        content = re.sub(r'^\s*```(?:markdown|md|text)?\s*$', '', content or '', flags=re.MULTILINE)
        content = content.strip()
        
        pdf.set_font("Helvetica", "", 11)
        for line in content.split('\n'):
            stripped = line.strip()
            try:
                self._render_line(pdf, stripped, safe, safe_multi_cell)
            except Exception:
                # Page space exhausted — add new page and retry
                try:
                    pdf.add_page()
                    pdf.set_font("Helvetica", "", 11)
                    self._render_line(pdf, stripped, safe, safe_multi_cell)
                except Exception:
                    # Truly unrenderable line — skip it
                    pdf.set_font("Helvetica", "", 11)
        
        return pdf.output()

    def create_docx(self, title: str, content: str) -> bytes:
        """Generate DOCX from markdown content."""
        doc = Document()
        
        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Strip code fences from content
        content = re.sub(r'^\s*```(?:markdown|md|text)?\s*$', '', content or '', flags=re.MULTILINE)
        content = content.strip()
        
        lines = content.split('\n')
        for line in lines:
            # Skip code fence lines
            if line.strip().startswith('```'):
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.strip():
                # Clean markdown formatting
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
                clean_line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_line)
                doc.add_paragraph(clean_line)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def create_document(self, title: str, content: str, format: str) -> dict:
        if format.lower() == "pdf":
            doc_bytes = self.create_pdf(title, content)
            mime_type = "application/pdf"
            filename = f"{title[:50]}.pdf"
        elif format.lower() == "docx":
            doc_bytes = self.create_docx(title, content)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{title[:50]}.docx"
        else:
            raise ValueError("Format must be 'pdf' or 'docx'")
        
        encoded = base64.b64encode(doc_bytes).decode('utf-8')
        return {"filename": filename, "mime_type": mime_type, "data": encoded}
